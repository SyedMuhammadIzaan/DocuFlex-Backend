from fastapi import APIRouter, UploadFile, Form
from fastapi.responses import StreamingResponse
import io
import pandas as pd
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


router = APIRouter(prefix="/convert", tags=["File Converter"])

@router.post("/")
async def convert_file(file: UploadFile, target_format: str = Form(...)):
    print(f"📦 Received file: {file.filename}, target: {target_format}")
    content = await file.read()
    input_stream = io.BytesIO(content)
    output_stream = io.BytesIO()

    name, ext = file.filename.rsplit(".", 1)
    ext = ext.lower()
    target_format = target_format.lower()
    print(f"➡️ Source extension: {ext}, Target: {target_format}")

    if ext == "txt" and target_format == "pdf":
        c = canvas.Canvas(output_stream, pagesize=letter)
        text = content.decode("utf-8").splitlines()
        width, height = letter
        x, y = 50, height - 50
        for line in text:
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(x, y, line.strip())
            y -= 15
        c.save()

    elif ext == "csv" and target_format == "xlsx":
        df = pd.read_csv(input_stream)
        df.to_excel(output_stream, index=False)

    elif ext in ["jpg", "jpeg", "png", "webp"] and target_format in ["png", "jpg", "webp"]:
        img = Image.open(input_stream)
        if target_format == "jpg":
            target_format = "JPEG"
        img.save(output_stream, format=target_format.upper())

    elif ext in ["jpg", "jpeg", "png", "webp"] and target_format == "pdf":
        # ✅ Convert image to PDF
        img = Image.open(input_stream).convert("RGB")
        img.save(output_stream, format="PDF")

    elif ext in ["jpg","jpeg","png","webp"] and target_format == "docx":
        img=Image.open(input_stream)
        extracted_text=pytesseract.image_to_string(img)
        
        doc=Document()
        doc.add_heading('Extracted Text from Image',level=1)
        doc.add_paragraph(extracted_text)
        doc_output=io.BytesIO()
        doc.save(doc_output)
        doc_output.seek(0)
        output_stream=doc_output
        
    else:
        print(f"❌ Unsupported conversion from {ext} to {target_format}")
        return {"error": f"Unsupported conversion from {ext} to {target_format}"}

    output_stream.seek(0)
    media_type = {
        "pdf": "application/pdf",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "webp": "image/webp",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }.get(target_format, "application/octet-stream")

    return StreamingResponse(
        output_stream,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={name}.{target_format}"}
    )
